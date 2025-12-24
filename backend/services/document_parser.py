"""
Multi-Format Document Parser.
Supports: PDF, Excel (.xlsx, .xls), CSV, Word (.docx)
"""

import os
import re
import json
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple

# File magic headers for detection
MAGIC_HEADERS = {
    b'%PDF': 'pdf',
    b'PK\x03\x04': 'xlsx_or_docx',  # ZIP-based formats
    b'\xd0\xcf\x11\xe0': 'xls',  # Old Excel format
}


def detect_file_type(file_path: str) -> str:
    """
    Detect file type using extension and magic header.
    Returns: 'pdf', 'xlsx', 'xls', 'csv', 'docx', 'unknown'
    """
    ext = Path(file_path).suffix.lower()
    
    # Try magic header first
    try:
        with open(file_path, 'rb') as f:
            header = f.read(8)
            
            for magic, ftype in MAGIC_HEADERS.items():
                if header.startswith(magic):
                    if ftype == 'xlsx_or_docx':
                        # Distinguish by extension
                        if ext in ['.xlsx', '.xls']:
                            return 'xlsx'
                        elif ext == '.docx':
                            return 'docx'
                        return 'xlsx'  # Default to xlsx for ZIP
                    return ftype
    except:
        pass
    
    # Fallback to extension
    ext_map = {
        '.pdf': 'pdf',
        '.xlsx': 'xlsx',
        '.xls': 'xls',
        '.csv': 'csv',
        '.docx': 'docx',
    }
    return ext_map.get(ext, 'unknown')


def parse_pdf_document(file_path: str) -> Dict[str, Any]:
    """
    Parse PDF using Docling (primary) or PyPDF2 (fallback).
    """
    text_blocks = []
    tables = []
    metadata = {}
    
    # Try Docling first
    try:
        from docling.document_converter import DocumentConverter
        converter = DocumentConverter()
        result = converter.convert(file_path)
        
        # Extract text
        if hasattr(result, 'document') and hasattr(result.document, 'export_to_markdown'):
            text = result.document.export_to_markdown()
            text_blocks.append(text)
        
        # Extract tables if available
        if hasattr(result, 'document') and hasattr(result.document, 'tables'):
            for table in result.document.tables:
                if hasattr(table, 'export_to_dataframe'):
                    df = table.export_to_dataframe()
                    tables.append(df.to_dict('records'))
        
        return {
            "text": "\n".join(text_blocks),
            "tables": tables,
            "meta": {"parser": "docling", "pages": getattr(result, 'num_pages', 1)},
            "document_type": "PDF",
        }
    except Exception as e:
        pass
    
    # Fallback to PyPDF2
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_blocks.append(text)
            metadata = {
                "parser": "pypdf2",
                "pages": len(reader.pages),
            }
        
        return {
            "text": "\n".join(text_blocks),
            "tables": tables,
            "meta": metadata,
            "document_type": "PDF",
        }
    except Exception as e:
        pass
    
    # Final fallback - just return empty
    return {
        "text": "",
        "tables": [],
        "meta": {"parser": "none", "error": "Could not parse PDF"},
        "document_type": "PDF",
    }


def parse_excel_document(file_path: str) -> Dict[str, Any]:
    """
    Parse Excel (.xlsx, .xls) using pandas/openpyxl.
    """
    text_blocks = []
    tables = []
    sheet_names = []
    
    try:
        import pandas as pd
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        for sheet_name in sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            
            # Convert to text format
            text_blocks.append(f"\n=== SHEET: {sheet_name} ===\n")
            
            # Add header row
            text_blocks.append(f"Columns: {', '.join(str(c) for c in df.columns)}\n")
            
            # Add rows as text
            for idx, row in df.iterrows():
                row_text = f"Row {idx + 1}: "
                row_parts = []
                for col, val in row.items():
                    if pd.notna(val):
                        row_parts.append(f"{col}={val}")
                row_text += ", ".join(row_parts)
                text_blocks.append(row_text)
            
            # Store as table grid
            table_data = df.fillna("").to_dict('records')
            tables.append({
                "sheet": sheet_name,
                "data": table_data,
                "columns": list(df.columns),
                "rows": len(df),
            })
            
            # Also add markdown-formatted table for better LLM parsing
            text_blocks.append(f"\n--- TABLE: {sheet_name} ---\n")
            # Add header row
            text_blocks.append("| " + " | ".join(str(c) for c in df.columns) + " |")
            text_blocks.append("| " + " | ".join(["---"] * len(df.columns)) + " |")
            # Add data rows (limit to 100 rows for token efficiency)
            for idx, row in df.head(100).iterrows():
                row_values = [str(val) if pd.notna(val) else "" for val in row]
                text_blocks.append("| " + " | ".join(row_values) + " |")
        
        return {
            "text": "\n".join(text_blocks),
            "tables": tables,
            "meta": {"parser": "pandas", "sheets": sheet_names},
            "document_type": "EXCEL",
        }
    except Exception as e:
        return {
            "text": "",
            "tables": [],
            "meta": {"parser": "none", "error": str(e)},
            "document_type": "EXCEL",
        }


def parse_csv_document(file_path: str) -> Dict[str, Any]:
    """
    Parse CSV using pandas.
    """
    text_blocks = []
    tables = []
    
    try:
        import pandas as pd
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                break
            except:
                continue
        else:
            df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
        
        # Convert to text
        text_blocks.append(f"=== CSV DATA ===\n")
        text_blocks.append(f"Columns: {', '.join(str(c) for c in df.columns)}\n")
        
        for idx, row in df.iterrows():
            row_text = f"Row {idx + 1}: "
            row_parts = []
            for col, val in row.items():
                if pd.notna(val):
                    row_parts.append(f"{col}={val}")
            row_text += ", ".join(row_parts)
            text_blocks.append(row_text)
        
        # Store as table
        table_data = df.fillna("").to_dict('records')
        tables.append({
            "sheet": "CSV",
            "data": table_data,
            "columns": list(df.columns),
            "rows": len(df),
        })
        
        # Also add markdown-formatted table for better LLM parsing
        text_blocks.append(f"\n--- TABLE: CSV Data ---\n")
        # Add header row
        text_blocks.append("| " + " | ".join(str(c) for c in df.columns) + " |")
        text_blocks.append("| " + " | ".join(["---"] * len(df.columns)) + " |")
        # Add data rows (limit to 100 rows for token efficiency)
        for idx, row in df.head(100).iterrows():
            row_values = [str(val) if pd.notna(val) else "" for val in row]
            text_blocks.append("| " + " | ".join(row_values) + " |")
        
        return {
            "text": "\n".join(text_blocks),
            "tables": tables,
            "meta": {"parser": "pandas", "rows": len(df)},
            "document_type": "CSV",
        }
    except Exception as e:
        return {
            "text": "",
            "tables": [],
            "meta": {"parser": "none", "error": str(e)},
            "document_type": "CSV",
        }


def parse_word_document(file_path: str) -> Dict[str, Any]:
    """
    Parse Word (.docx) using python-docx.
    """
    text_blocks = []
    tables = []
    
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_blocks.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = row_data
                else:
                    if headers:
                        table_data.append(dict(zip(headers, row_data)))
                    else:
                        table_data.append({"row": row_idx, "cells": row_data})
            
            tables.append({
                "table_index": table_idx,
                "data": table_data,
                "columns": headers,
                "rows": len(table_data),
            })
            
            # Also add table as markdown for better LLM parsing
            text_blocks.append(f"\n--- TABLE {table_idx + 1} ---\n")
            if headers:
                text_blocks.append("| " + " | ".join(headers) + " |")
                text_blocks.append("| " + " | ".join(["---"] * len(headers)) + " |")
                for row_data in table_data[:100]:  # Limit to 100 rows
                    if isinstance(row_data, dict):
                        row_values = [str(row_data.get(h, "")) for h in headers]
                        text_blocks.append("| " + " | ".join(row_values) + " |")
            else:
                # Fallback for tables without headers
                for row in table.rows[:100]:
                    text_blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        
        return {
            "text": "\n".join(text_blocks),
            "tables": tables,
            "meta": {"parser": "python-docx", "paragraphs": len(doc.paragraphs)},
            "document_type": "WORD",
        }
    except Exception as e:
        return {
            "text": "",
            "tables": [],
            "meta": {"parser": "none", "error": str(e)},
            "document_type": "WORD",
        }


def parse_document(file_path: str) -> Dict[str, Any]:
    """
    Master document parser. Detects file type and routes to appropriate parser.
    
    Returns:
        {
            "text": "...",
            "tables": [...],
            "meta": {...},
            "document_type": "PDF|EXCEL|CSV|WORD"
        }
    """
    if not os.path.exists(file_path):
        return {
            "text": "",
            "tables": [],
            "meta": {"error": "File not found"},
            "document_type": "UNKNOWN",
        }
    
    file_type = detect_file_type(file_path)
    
    parsers = {
        'pdf': parse_pdf_document,
        'xlsx': parse_excel_document,
        'xls': parse_excel_document,
        'csv': parse_csv_document,
        'docx': parse_word_document,
    }
    
    parser = parsers.get(file_type)
    if parser:
        result = parser(file_path)
        return result
    else:
        return {
            "text": "",
            "tables": [],
            "meta": {"error": f"Unsupported file type: {file_type}"},
            "document_type": "UNKNOWN",
        }


def merge_parsed_documents(parsed_docs: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Merge multiple parsed documents into a single context for LLM extraction.
    """
    all_text = []
    all_tables = []
    
    for i, doc in enumerate(parsed_docs):
        doc_type = doc.get("document_type", "UNKNOWN")
        text = doc.get("text", "")
        tables = doc.get("tables", [])
        
        # Add document header
        all_text.append(f"\n\n========== DOCUMENT {i + 1} ({doc_type}) ==========\n")
        all_text.append(text)
        
        # Collect tables
        for table in tables:
            all_tables.append({
                "source_document": i + 1,
                "document_type": doc_type,
                **table
            })
    
    return {
        "full_context_text": "\n".join(all_text),
        "tables": all_tables,
        "document_count": len(parsed_docs),
    }
